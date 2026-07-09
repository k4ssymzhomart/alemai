"""Low-level python-docx helpers shared by the Epic D document templates.

No jinja placeholders: documents are assembled programmatically so a rendered
file can never leak a ``{{ }}`` token (asserted in tests). Kazakh strings are
flagged NEEDS-NATIVE-REVIEW inline (a trailing marker paragraph) and mirrored in
docs/NATIVE-REVIEW-QUEUE.md.
"""

from __future__ import annotations

import io
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

# Visible marker so an un-reviewed kk artifact is obvious on paper and in tests.
NATIVE_REVIEW_MARK = "[kk: NEEDS-NATIVE-REVIEW]"
ORG_HEADER_RU = "ГП №14"
ORG_HEADER_KK = "№14 қалалық емхана"
# Brand letterhead (docs/25 H1). __file__-relative so it resolves both in the
# api container (/app/...) and in local pytest runs.
_LOGO_PATH = Path(__file__).resolve().parent / "assets" / "logo.png"


def new_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    return doc


def add_org_header(doc: Document, lang: str, date_str: str) -> None:
    """Logo letterhead + org line + date, like an official document header."""
    if _LOGO_PATH.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.add_run().add_picture(str(_LOGO_PATH), width=Mm(42))
    org = ORG_HEADER_KK if lang == "kk" else ORG_HEADER_RU
    p = doc.add_paragraph()
    run = p.add_run(org)
    run.bold = True
    run.font.size = Pt(12)
    date_label = "күні" if lang == "kk" else "дата"
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dp.add_run(f"{date_label}: {date_str}")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_kv_table(doc: Document, rows: Iterable[tuple[str, str]]) -> None:
    """Two-column key/value table (label | value), thin grid."""
    rows = list(rows)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = ""
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].text = value


def add_citation_block(doc: Document, lang: str, label: str, snippet: str) -> None:
    """Render the legal basis: «Основание/Дереккөз» + label + quoted snippet."""
    heading = "Дереккөз (нормативтік негіз):" if lang == "kk" else "Основание (НПА):"
    add_heading(doc, heading)
    p = doc.add_paragraph()
    p.add_run(label).italic = True
    q = doc.add_paragraph()
    q.paragraph_format.left_indent = Pt(18)
    q.add_run(f"«{snippet}»").font.size = Pt(10)


def add_signature_block(doc: Document, lang: str) -> None:
    doc.add_paragraph()
    if lang == "kk":
        doc.add_paragraph("Бас дәрігер ________________________  (қолы, Т.А.Ә.)")
        doc.add_paragraph("М.О.")
    else:
        doc.add_paragraph("Главный врач ________________________  (подпись, Ф.И.О.)")
        doc.add_paragraph("М.П.")


def add_native_review_footer(doc: Document, lang: str) -> None:
    """For kk artifacts: an explicit NEEDS-NATIVE-REVIEW marker (no demo badge)."""
    if lang != "kk":
        return
    doc.add_paragraph()
    mp = doc.add_paragraph()
    mr = mp.add_run(NATIVE_REVIEW_MARK)
    mr.font.size = Pt(8)
    mr.italic = True


def render(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
